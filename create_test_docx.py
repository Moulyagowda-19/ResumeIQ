import docx

doc = docx.Document()
doc.add_heading('Test Resume', 0)
doc.add_paragraph('Name: Test User')
doc.add_paragraph('Email: test@example.com')
doc.add_paragraph('Skills: Python, Flask, SQL')
doc.add_paragraph('Experience: 5 years of Python development.')

doc.save('test_resume.docx')
print("test_resume.docx created successfully.")
